"""Report assembler — convert agent markdown to a formatted .docx using a reference template."""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt

from engineering_hub.templates.models import ReportSkeleton

logger = logging.getLogger(__name__)

_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"\*(.+?)\*")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")
_LIST_ITEM_RE = re.compile(r"^(\s*)[-*•]\s+(.+)$")
_NUMBERED_ITEM_RE = re.compile(r"^(\s*)\d+[.)]\s+(.+)$")
_BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")
_PLACEHOLDER_RE = re.compile(r"\[INSERT:\s*([^\]]+)\]")


class ReportAssembler:
    """Assemble a formatted .docx from markdown content and a reference template."""

    def __init__(self, skeleton: ReportSkeleton) -> None:
        self.skeleton = skeleton
        self._style_map: dict[str, str] = {}

    def assemble(self, markdown: str, output_path: Path) -> Path:
        """Convert markdown to a .docx file styled by the reference template.

        Args:
            markdown: The agent-produced markdown content.
            output_path: Where to write the final .docx.

        Returns:
            The output path written.
        """
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        ref_path = Path(self.skeleton.reference_docx_path)
        if ref_path.exists():
            doc = Document(str(ref_path))
            self._clear_body(doc)
        else:
            logger.warning(
                "Reference docx not found at %s; creating blank document", ref_path
            )
            doc = Document()

        self._build_style_map(doc)
        self._populate_body(doc, markdown)

        doc.save(str(output_path))
        logger.info("Assembled report saved: %s", output_path)
        return output_path

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _clear_body(doc: Document) -> None:
        """Remove all body paragraphs and tables, preserving headers/footers/styles."""
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("p", "tbl"):
                body.remove(child)

    def _build_style_map(self, doc: Document) -> None:
        """Map markdown heading levels to available DOCX style names."""
        available = {s.name for s in doc.styles if s.name}

        for level in range(1, 7):
            candidates = [
                f"Heading {level}",
                f"heading {level}",
                f"Heading{level}",
            ]
            for c in candidates:
                if c in available:
                    self._style_map[f"h{level}"] = c
                    break

        if "Normal" in available:
            self._style_map["body"] = "Normal"
        if "List Bullet" in available:
            self._style_map["list_bullet"] = "List Bullet"
        if "List Number" in available:
            self._style_map["list_number"] = "List Number"
        if "Quote" in available:
            self._style_map["quote"] = "Quote"
        elif "Block Text" in available:
            self._style_map["quote"] = "Block Text"

    def _populate_body(self, doc: Document, markdown: str) -> None:
        """Parse markdown and add elements to the document body."""
        lines = markdown.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # Blank line
            if not line.strip():
                i += 1
                continue

            # Heading
            m = _HEADING_RE.match(line)
            if m:
                level = len(m.group(1))
                text = m.group(2).strip()
                self._add_heading(doc, text, level)
                i += 1
                continue

            # Table (look ahead for separator)
            if (
                _TABLE_ROW_RE.match(line)
                and i + 1 < len(lines)
                and _TABLE_SEP_RE.match(lines[i + 1])
            ):
                i = self._add_table(doc, lines, i)
                continue

            # Blockquote
            m = _BLOCKQUOTE_RE.match(line)
            if m:
                quote_lines: list[str] = []
                while i < len(lines):
                    bq = _BLOCKQUOTE_RE.match(lines[i])
                    if bq:
                        quote_lines.append(bq.group(1))
                        i += 1
                    else:
                        break
                self._add_blockquote(doc, "\n".join(quote_lines))
                continue

            # List item (bullet)
            m = _LIST_ITEM_RE.match(line)
            if m:
                self._add_list_item(doc, m.group(2), bullet=True)
                i += 1
                continue

            # List item (numbered)
            m = _NUMBERED_ITEM_RE.match(line)
            if m:
                self._add_list_item(doc, m.group(2), bullet=False)
                i += 1
                continue

            # Regular paragraph
            self._add_paragraph(doc, line)
            i += 1

    def _add_heading(self, doc: Document, text: str, level: int) -> None:
        style = self._style_map.get(f"h{level}")
        if style:
            doc.add_paragraph(text, style=style)
        else:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.bold = True
                if run.font.size is None:
                    run.font.size = Pt(max(18 - (level * 2), 10))

    def _add_paragraph(self, doc: Document, text: str) -> None:
        style = self._style_map.get("body", "Normal")
        p = doc.add_paragraph(style=style)
        self._add_rich_text(p, text)

    def _add_rich_text(self, paragraph: Any, text: str) -> None:
        """Parse inline markdown (bold, italic, placeholders) into runs."""
        parts = self._split_inline(text)
        for content, bold, italic in parts:
            run = paragraph.add_run(content)
            if bold:
                run.bold = True
            if italic:
                run.italic = True

    @staticmethod
    def _split_inline(text: str) -> list[tuple[str, bool, bool]]:
        """Split text into (content, bold, italic) segments."""
        segments: list[tuple[str, bool, bool]] = []
        pos = 0
        combined = re.compile(r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)")

        for m in combined.finditer(text):
            if m.start() > pos:
                segments.append((text[pos : m.start()], False, False))

            if m.group(2):  # ***bold italic***
                segments.append((m.group(2), True, True))
            elif m.group(3):  # **bold**
                segments.append((m.group(3), True, False))
            elif m.group(4):  # *italic*
                segments.append((m.group(4), False, True))
            pos = m.end()

        if pos < len(text):
            segments.append((text[pos:], False, False))

        return segments if segments else [(text, False, False)]

    def _add_list_item(self, doc: Document, text: str, bullet: bool = True) -> None:
        if bullet:
            style = self._style_map.get("list_bullet", "List Bullet")
        else:
            style = self._style_map.get("list_number", "List Number")

        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
            text = f"{'•' if bullet else '-'} {text}"

        self._add_rich_text(p, text)

    def _add_blockquote(self, doc: Document, text: str) -> None:
        style = self._style_map.get("quote")
        if style:
            try:
                p = doc.add_paragraph(text, style=style)
                return
            except KeyError:
                pass
        p = doc.add_paragraph(text)
        pf = p.paragraph_format
        pf.left_indent = Inches(0.5)
        for run in p.runs:
            run.italic = True

    def _add_table(self, doc: Document, lines: list[str], start: int) -> int:
        """Parse a markdown table starting at `start` and add it to the doc.

        Returns the next line index after the table.
        """
        header_line = lines[start]
        headers = [c.strip() for c in header_line.strip("|").split("|")]
        i = start + 2  # skip header + separator

        rows: list[list[str]] = []
        while i < len(lines) and _TABLE_ROW_RE.match(lines[i]):
            cells = [c.strip() for c in lines[i].strip("|").split("|")]
            rows.append(cells)
            i += 1

        ncols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=ncols)
        table.style = "Table Grid"

        # Header row
        for j, h in enumerate(headers):
            if j < ncols:
                cell = table.rows[0].cells[j]
                cell.text = h
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True

        # Data rows
        for r_idx, row_data in enumerate(rows):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < ncols:
                    table.rows[r_idx + 1].cells[c_idx].text = cell_text

        return i
