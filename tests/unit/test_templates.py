"""Tests for the report template analysis, assembly, and org-context pipeline."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches, Pt

from engineering_hub.templates.models import (
    HeaderFooterSpec,
    ReportSkeleton,
    SectionPattern,
    StyleSpec,
    TablePattern,
)


# ---------------------------------------------------------------------------
# Fixtures: create sample .docx files for analysis
# ---------------------------------------------------------------------------


def _make_sample_report(path: Path, title: str, include_table: bool = True) -> None:
    """Create a minimal .docx report with headings, body, and optionally a table."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = "Acme Consulting — Confidential"

    doc.add_heading(title, level=1)
    doc.add_paragraph(
        "This report presents the findings of the acoustic assessment "
        "conducted at the project site."
    )
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("The purpose of this document is to outline test procedures.")
    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(
        "Testing was performed in accordance with ASTM E336-17a. "
        "All equipment was calibrated prior to use."
    )
    doc.add_heading("Results", level=2)
    if include_table:
        table = doc.add_table(rows=3, cols=3)
        table.rows[0].cells[0].text = "Location"
        table.rows[0].cells[1].text = "STC Rating"
        table.rows[0].cells[2].text = "Pass/Fail"
        table.rows[1].cells[0].text = "Suite 101"
        table.rows[1].cells[1].text = "52"
        table.rows[1].cells[2].text = "Pass"
        table.rows[2].cells[0].text = "Suite 102"
        table.rows[2].cells[1].text = "48"
        table.rows[2].cells[2].text = "Fail"
    doc.add_heading("Conclusions", level=2)
    doc.add_paragraph("Based on the findings, the following recommendations are made.")
    doc.add_heading("References", level=2)
    doc.add_paragraph("ASTM E336-17a Standard Test Method for Measurement.")

    doc.save(str(path))


@pytest.fixture
def docx_corpus(tmp_path: Path) -> Path:
    """Create a temporary directory with several sample .docx files."""
    corpus_dir = tmp_path / "reports"
    corpus_dir.mkdir()

    for i in range(5):
        _make_sample_report(
            corpus_dir / f"report_{i}.docx",
            title=f"Acoustic Assessment Report #{i + 1}",
            include_table=(i < 4),
        )
    return corpus_dir


@pytest.fixture
def sample_skeleton(tmp_path: Path) -> ReportSkeleton:
    """Create a minimal ReportSkeleton for unit tests."""
    return ReportSkeleton(
        name="Test Template",
        source_doc_count=5,
        sections=[
            SectionPattern(heading="Introduction", level=2, frequency=1.0, typical_content_type="prose"),
            SectionPattern(heading="Methodology", level=2, frequency=1.0, typical_content_type="prose"),
            SectionPattern(heading="Results", level=2, frequency=0.8, typical_content_type="table"),
            SectionPattern(heading="Conclusions", level=2, frequency=1.0, typical_content_type="prose"),
            SectionPattern(heading="References", level=2, frequency=1.0, typical_content_type="list"),
            SectionPattern(heading="Appendix A", level=2, frequency=0.3, typical_content_type="prose"),
        ],
        styles={
            "Normal": StyleSpec(name="Normal", font_name="Calibri", font_size_pt=11.0),
            "Heading 1": StyleSpec(name="Heading 1", font_name="Calibri", font_size_pt=16.0, bold=True),
        },
        header=HeaderFooterSpec(text="Acme Consulting"),
        footer=HeaderFooterSpec(has_page_numbers=True),
        page_margins_inches={"top": 1.0, "bottom": 1.0, "left": 1.25, "right": 1.25},
        table_patterns=[
            TablePattern(
                label="Location / STC Rating / Pass/Fail",
                column_headers=["Location", "STC Rating", "Pass/Fail"],
                frequency=0.8,
            )
        ],
        reference_docx_path="",
    )


# ---------------------------------------------------------------------------
# Test: Models
# ---------------------------------------------------------------------------


class TestReportSkeletonModel:
    def test_save_and_load(self, tmp_path: Path, sample_skeleton: ReportSkeleton) -> None:
        path = tmp_path / "out" / "skeleton.json"
        sample_skeleton.save(path)

        assert path.exists()
        loaded = ReportSkeleton.load(path)
        assert loaded.name == sample_skeleton.name
        assert loaded.source_doc_count == 5
        assert len(loaded.sections) == 6
        assert loaded.sections[0].heading == "Introduction"

    def test_format_for_agent(self, sample_skeleton: ReportSkeleton) -> None:
        text = sample_skeleton.format_for_agent()
        assert "Report Template: Test Template" in text
        assert "Required Sections" in text
        assert "Introduction" in text
        assert "Methodology" in text
        assert "Optional Sections" in text
        assert "Appendix A" in text
        assert "Common Table Structures" in text
        assert "Location" in text

    def test_frequency_threshold(self, sample_skeleton: ReportSkeleton) -> None:
        text = sample_skeleton.format_for_agent()
        lines = text.split("\n")
        required_block = []
        optional_block = []
        in_required = False
        in_optional = False
        for line in lines:
            if "### Required Sections" in line:
                in_required = True
                in_optional = False
                continue
            if "### Optional Sections" in line:
                in_required = False
                in_optional = True
                continue
            if line.startswith("###"):
                in_required = False
                in_optional = False
            if in_required and line.startswith("- "):
                required_block.append(line)
            if in_optional and line.startswith("- "):
                optional_block.append(line)

        assert len(required_block) == 5
        assert len(optional_block) == 1
        assert "Appendix A" in optional_block[0]


# ---------------------------------------------------------------------------
# Test: TemplateAnalyzer
# ---------------------------------------------------------------------------


class TestTemplateAnalyzer:
    def test_analyze_produces_skeleton(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        skeleton = analyzer.analyze(output_dir)

        assert skeleton.name == "Field Report"
        assert skeleton.source_doc_count == 5
        assert len(skeleton.sections) > 0
        assert any(s.heading == "Introduction" for s in skeleton.sections)
        assert any(s.heading == "Methodology" for s in skeleton.sections)

    def test_analyze_creates_files(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        analyzer.analyze(output_dir)

        assert (output_dir / "skeleton.json").exists()
        assert (output_dir / "reference.docx").exists()

    def test_analyze_extracts_styles(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        skeleton = analyzer.analyze(output_dir)

        assert len(skeleton.styles) > 0
        assert "Normal" in skeleton.styles

    def test_analyze_extracts_margins(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        skeleton = analyzer.analyze(output_dir)

        assert "top" in skeleton.page_margins_inches
        assert 0.5 < skeleton.page_margins_inches["top"] < 2.0

    def test_analyze_extracts_table_patterns(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        skeleton = analyzer.analyze(output_dir)

        assert len(skeleton.table_patterns) > 0
        first = skeleton.table_patterns[0]
        assert "Location" in first.column_headers

    def test_analyze_extracts_header(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        skeleton = analyzer.analyze(output_dir)

        assert skeleton.header.text is not None
        assert "Acme" in skeleton.header.text

    def test_analyze_empty_dir_raises(self, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        empty_dir = tmp_path / "empty"
        empty_dir.mkdir()
        analyzer = TemplateAnalyzer(empty_dir)
        with pytest.raises(FileNotFoundError, match="No .docx files found"):
            analyzer.analyze(tmp_path / "out")

    def test_reference_docx_is_valid(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        output_dir = tmp_path / "template_out"
        analyzer = TemplateAnalyzer(docx_corpus, name="Field Report")
        analyzer.analyze(output_dir)

        ref = Document(str(output_dir / "reference.docx"))
        headings = [
            p.text for p in ref.paragraphs
            if p.style.name and p.style.name.startswith("Heading")
        ]
        assert len(headings) > 0
        assert len(ref.tables) == 0


# ---------------------------------------------------------------------------
# Test: ReportAssembler
# ---------------------------------------------------------------------------


class TestReportAssembler:
    def test_assemble_basic(self, tmp_path: Path, sample_skeleton: ReportSkeleton) -> None:
        from engineering_hub.templates.assembler import ReportAssembler

        ref_path = tmp_path / "reference.docx"
        Document().save(str(ref_path))
        sample_skeleton.reference_docx_path = str(ref_path)

        assembler = ReportAssembler(sample_skeleton)
        md = (
            "# Acoustic Assessment Report\n\n"
            "## Introduction\n\n"
            "This report details the assessment.\n\n"
            "## Methodology\n\n"
            "Testing per **ASTM E336-17a**.\n\n"
            "## Results\n\n"
            "| Location | STC | Pass |\n"
            "|----------|-----|------|\n"
            "| Suite 101 | 52 | Yes |\n"
            "| Suite 102 | 48 | No |\n\n"
            "## Conclusions\n\n"
            "- Recommendation 1\n"
            "- Recommendation 2\n"
        )
        output_path = tmp_path / "output.docx"
        result = assembler.assemble(md, output_path)

        assert result == output_path
        assert output_path.exists()

        doc = Document(str(output_path))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        assert any("Introduction" in t for t in texts)
        assert any("Methodology" in t for t in texts)
        assert len(doc.tables) == 1

    def test_assemble_inline_formatting(self, tmp_path: Path, sample_skeleton: ReportSkeleton) -> None:
        from engineering_hub.templates.assembler import ReportAssembler

        ref_path = tmp_path / "reference.docx"
        Document().save(str(ref_path))
        sample_skeleton.reference_docx_path = str(ref_path)

        assembler = ReportAssembler(sample_skeleton)
        md = "## Test\n\nThis is **bold** and *italic* text.\n"
        output_path = tmp_path / "format_test.docx"
        assembler.assemble(md, output_path)

        doc = Document(str(output_path))
        body_paras = [p for p in doc.paragraphs if "bold" in p.text or "italic" in p.text]
        assert len(body_paras) == 1
        runs = body_paras[0].runs
        bold_runs = [r for r in runs if r.bold]
        italic_runs = [r for r in runs if r.italic]
        assert len(bold_runs) >= 1
        assert len(italic_runs) >= 1

    def test_assemble_blockquote(self, tmp_path: Path, sample_skeleton: ReportSkeleton) -> None:
        from engineering_hub.templates.assembler import ReportAssembler

        ref_path = tmp_path / "reference.docx"
        Document().save(str(ref_path))
        sample_skeleton.reference_docx_path = str(ref_path)

        assembler = ReportAssembler(sample_skeleton)
        md = "## Quote Test\n\n> This is a blockquote line.\n> Continued here.\n"
        output_path = tmp_path / "quote_test.docx"
        assembler.assemble(md, output_path)

        doc = Document(str(output_path))
        quote_paras = [p for p in doc.paragraphs if "blockquote" in p.text.lower() or "Continued" in p.text]
        assert len(quote_paras) >= 1

    def test_assemble_missing_reference(self, tmp_path: Path, sample_skeleton: ReportSkeleton) -> None:
        from engineering_hub.templates.assembler import ReportAssembler

        sample_skeleton.reference_docx_path = str(tmp_path / "nonexistent.docx")
        assembler = ReportAssembler(sample_skeleton)
        md = "## Test\n\nHello world.\n"
        output_path = tmp_path / "fallback.docx"
        assembler.assemble(md, output_path)

        assert output_path.exists()
        doc = Document(str(output_path))
        assert any("Hello world" in p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Test: Org Context Parser
# ---------------------------------------------------------------------------


class TestOrgContextParser:
    def test_parse_basic_note(self, tmp_path: Path) -> None:
        from engineering_hub.templates.org_context import parse_org_note

        note = tmp_path / "project.org"
        note.write_text(
            "#+title: Office Building Assessment\n"
            "#+filetags: :acoustic:consulting:\n"
            ":PROPERTIES:\n"
            ":CLIENT: Acme Construction\n"
            ":PROJECT_ID: 42\n"
            ":BUDGET: 45000\n"
            ":TECHNICAL_LEVEL: moderate\n"
            ":END:\n"
            "\n"
            "This project involves comprehensive acoustic testing.\n"
            "\n"
            "* Scope\n"
            "** ASTC testing per ASTM E336-17a\n"
            "** AIIC testing per ASTM E1007-16\n"
            "* Equipment\n"
            "- Sound level meter\n"
            "- Tapping machine\n",
            encoding="utf-8",
        )
        ctx = parse_org_note(note)

        assert ctx.project.title == "Office Building Assessment"
        assert ctx.project.client_name == "Acme Construction"
        assert ctx.project.id == 42
        assert ctx.project.budget == "45000"
        assert ctx.metadata["client_technical_level"] == "moderate"
        assert "acoustic" in ctx.metadata["filetags"]
        assert len(ctx.scope) > 0
        assert len(ctx.standards) >= 2

    def test_parse_missing_file_raises(self, tmp_path: Path) -> None:
        from engineering_hub.templates.org_context import parse_org_note

        with pytest.raises(FileNotFoundError):
            parse_org_note(tmp_path / "nonexistent.org")

    def test_parse_minimal_note(self, tmp_path: Path) -> None:
        from engineering_hub.templates.org_context import parse_org_note

        note = tmp_path / "minimal.org"
        note.write_text("#+title: Quick Note\n\nSome text here.\n")
        ctx = parse_org_note(note)

        assert ctx.project.title == "Quick Note"
        assert ctx.project.client_name == "Unknown Client"
        assert ctx.project.id == 0

    def test_standards_extraction(self, tmp_path: Path) -> None:
        from engineering_hub.templates.org_context import parse_org_note

        note = tmp_path / "standards.org"
        note.write_text(
            "#+title: Standards Test\n\n"
            "Testing per ASTM E336-17a and ISO 717-1.\n"
            "Also see ANSI S12.60.\n"
        )
        ctx = parse_org_note(note)

        ids = {s.id for s in ctx.standards}
        assert "ASTM E336-17a" in ids
        assert any("ISO" in s for s in ids)
        assert any("ANSI" in s for s in ids)


# ---------------------------------------------------------------------------
# Test: Full pipeline (analyzer -> skeleton -> assembler)
# ---------------------------------------------------------------------------


class TestFullPipeline:
    def test_analyze_then_assemble(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer
        from engineering_hub.templates.assembler import ReportAssembler

        template_dir = tmp_path / "templates"
        analyzer = TemplateAnalyzer(docx_corpus, name="Integration Test")
        skeleton = analyzer.analyze(template_dir)

        md_content = (
            "# Test Report\n\n"
            "## Introduction\n\n"
            "Project overview paragraph.\n\n"
            "## Methodology\n\n"
            "Testing per **ASTM E336-17a**, Section 7.\n\n"
            "## Results\n\n"
            "| Location | STC Rating | Pass/Fail |\n"
            "|----------|------------|----------|\n"
            "| Room A | 55 | Pass |\n\n"
            "## Conclusions\n\n"
            "All tested assemblies meet requirements.\n\n"
            "## References\n\n"
            "- ASTM E336-17a\n"
        )

        assembler = ReportAssembler(skeleton)
        output = tmp_path / "final_report.docx"
        assembler.assemble(md_content, output)

        assert output.exists()
        doc = Document(str(output))
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) == 1

        headings = [
            p.text for p in doc.paragraphs
            if p.style and p.style.name and p.style.name.startswith("Heading")
        ]
        assert "Introduction" in headings
        assert "Methodology" in headings
        assert "Results" in headings
        assert "Conclusions" in headings

    def test_skeleton_roundtrip(self, docx_corpus: Path, tmp_path: Path) -> None:
        from engineering_hub.templates.analyzer import TemplateAnalyzer

        template_dir = tmp_path / "templates"
        analyzer = TemplateAnalyzer(docx_corpus, name="Roundtrip Test")
        original = analyzer.analyze(template_dir)

        loaded = ReportSkeleton.load(template_dir / "skeleton.json")

        assert loaded.name == original.name
        assert loaded.source_doc_count == original.source_doc_count
        assert len(loaded.sections) == len(original.sections)
        assert len(loaded.styles) == len(original.styles)
        assert loaded.page_margins_inches == original.page_margins_inches

    def test_org_context_with_template(self, docx_corpus: Path, tmp_path: Path) -> None:
        """End-to-end: org note + skeleton -> formatted context for agent."""
        from engineering_hub.context.formatters import ContextFormatter
        from engineering_hub.core.constants import AgentType
        from engineering_hub.templates.analyzer import TemplateAnalyzer
        from engineering_hub.templates.org_context import parse_org_note

        template_dir = tmp_path / "templates"
        analyzer = TemplateAnalyzer(docx_corpus, name="E2E Test")
        skeleton = analyzer.analyze(template_dir)

        note = tmp_path / "project.org"
        note.write_text(
            "#+title: Test Project\n"
            ":PROPERTIES:\n"
            ":CLIENT: Test Corp\n"
            ":END:\n"
            "\n"
            "* Field Testing\n"
            "Per ASTM E336-17a.\n"
        )
        ctx = parse_org_note(note)
        ctx.metadata["template_skeleton_block"] = skeleton.format_for_agent()

        formatted = ContextFormatter.format(ctx, AgentType.TECHNICAL_WRITER)
        assert "Report Template: E2E Test" in formatted
        assert "Required Sections" in formatted
        assert "Test Corp" in formatted
